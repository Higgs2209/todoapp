import tkinter as tk
from tkinter import filedialog
import customtkinter as ctk
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from tkinter import StringVar, TOP
from tkinterdnd2 import TkinterDnD, DND_ALL
import os
import pathlib
from pathlib import Path
import pyuac
from striprtf.striprtf import rtf_to_text
from docx import Document
from docx.shared import Inches
import sentencepiece
from tqdm import tqdm
import docx2txt 
from pypdf import PdfReader
import time
import torch
import psutil
from torch.quantization import quantize_dynamic
import ttkbootstrap as ttk
import ttkbootstrap.constants  
import threading
from ttkbootstrap.tooltip import ToolTip


#run as administrator
if __name__ == "__main__":
    if not pyuac.isUserAdmin():
        print("Re-launching as admin!")
        pyuac.runAsAdmin()
    else:        
        pass# Already an admin here.


#CPU System Settings
device = "cuda:0" if torch.cuda.is_available() else "cpu"
for p in psutil.process_iter():
     if p.name() == "Case_Summariser.exe":
          p.nice(psutil.BELOW_NORMAL_PRIORITY_CLASS)
          
#Select file Path
def file_path():
    global filePath
    global file_name 
    filePath = filedialog.askopenfilename()
    global outputPathText
    file_name = Path(filePath).stem
    FileToString()
  
          
def savedfileName():
    global SavingFileName
    SavingFileName = saving_text.get()
    print(SavingFileName)
       
# Create a global variable to keep track of the label
outputPathText = None
                  
#clear output path
def clear_output_path():
    global outputPathText
    if outputPathText:
        outputPathText.pack_forget()
        outputPathText = None

#Select Saving Location 
def SavingLocation():
    clear_output_path()
    global output_path
    output_path = ""
    output_path = filedialog.askdirectory()
    global outputPathText
    outputPathText = ttk.Label(root, text="Saving Location:" + output_path)
    outputPathText.pack()
        
    output_path = str(os.path.join(Path.home(), "Downloads"))
    

    
    
          

   



def FileToString():
    loadingDoc = ttk.Label(root, text="Loading Document...")
    loadingDoc.pack()
    root.update()
    print("Test Loading Labels...")
    global text
    filename = filePath
    extension = os.path.splitext(filename)[1]
    print(extension)
    global summarise
    if (extension == ".rtf"):
        with open(filePath, 'r') as file: 
            rtf_text = file.read() 
            text = rtf_to_text(rtf_text) 
            print(text)
            summarise = ttk.Button(root, text="Summarise Case", command=threading.Thread(target=Summarise).start)
            summarise.pack(padx=10, pady=10)
            #open file in word
            
    elif (extension == ".docx"):
        text = docx2txt.process(filePath)
        print(text)
        summarise = ttk.Button(root, text="Summarise Case", command=threading.Thread(target=Summarise).start)
        summarise.pack(padx=10, pady=10)
        #outputPathText.pack_forget()        
    elif (extension == ".txt"):
        file = open(filePath, 'r') 
        text = file.read()
        file.close()   
        print(text)
        summarise = ttk.Button(root, text="Summarise Case", command=threading.Thread(target=Summarise).start)
        summarise.pack(padx=10, pady=10)
        #outputPathText.pack_forget()
    elif (extension == ".pdf"):
        reader = PdfReader(filePath)
        text = ""
        print("Loading Document...")
        for page in reader.pages:
            text += page.extract_text() + "\n"
        global DocLoaded
        DocLoaded = ttk.Label(root, text="Document Loaded")
        DocLoaded.pack()
        loadingDoc.pack_forget()
        print("Document Loaded")
        #summarise button
        summarise = ttk.Button(root, text="Summarise Case", command=threading.Thread(target=Summarise).start)
        summarise.pack(padx=10, pady=10)
        

        
        


#save summary as a new file

def DocGen():
    document = Document()

    document.add_heading("summary", 0)

    p = document.add_paragraph(summary)
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.save(output_path + "/" + SavingFileName + ".docx")
    print(output_path + "/" + SavingFileName + ".docx")

#Text Summarising Code
def Summarise():
    global progress_bar
    progress_bar = ttk.Progressbar(root, orient='horizontal', mode='determinate', length=300, value=0)
    progress_bar.pack(pady=20)
    for p in psutil.process_iter():
     if p.name() == "Case_Summariser.exe":
          p.nice(psutil.BELOW_NORMAL_PRIORITY_CLASS)
    Summarising_Case = ttk.Label(root, text="Summarising Case")
    Summarising_Case.pack()
    root.update()
    print("Summarising Case")
    #time.sleep(10)
    global summary
    global loadingSummaryText
    loadingSummaryText=ttk.Label(root, text="Loading Tokens...")
    loadingSummaryText.pack()
    input_tokenized = tokenizer.encode(text, return_tensors='pt',max_length=1024,truncation=True)#.to(device)
    progress_bar['value'] += 25
    loadingSummaryText.pack_forget()
    loadingSummaryText=ttk.Label(root, text="Loading Summary IDs...")
    loadingSummaryText.pack()
    summary_ids = tqdm(model_quantized.generate(input_tokenized,
                                    num_beams=4,
                                    no_repeat_ngram_size=3,
                                    length_penalty=2.0,
                                    min_length=150,
                                    max_length=999999999999999999,
                                    early_stopping=True).to(device))
    progress_bar['value'] += 50
    print("Model Generated")
    progress_bar['value'] += 75
    loadingSummaryText.pack_forget()
    loadingSummaryText=ttk.Label(root, text="Summary")
    loadingSummaryText.pack()
    summary = tqdm([tokenizer.decode(g, skip_special_tokens=True, clean_up_tokenization_spaces=False) for g in summary_ids][0])
    print("Summary Generated")
    DocGen()
    print(summary)
    print("Summarised Case")
    progress_bar['value'] += 100
    global case_summarised
    Summarising_Case.pack_forget()
    case_summarised = ttk.Label(root, text="Case Summarised")
    case_summarised.pack()
    global open_summary_button
    open_summary_button = ttk.Button(root, text="Open Summary", command = open_summary)
    open_summary_button.pack(pady=10)
    ToolTip(open_summary_button, text="Open the summarised file in Word")
          

#threading        
print("Loading Model...")
def load_tokenizer():
    global tokenizer
    tokenizer = AutoTokenizer.from_pretrained("nsi319/legal-pegasus")
    print("Task 1/3 Completed!")
def load_model():
    global model
    model = AutoModelForSeq2SeqLM.from_pretrained("nsi319/legal-pegasus")
    print("Task 2/3 Completed!")

def load_model_quantized():
    global model_quantized
    model_quantized = quantize_dynamic(model, dtype=torch.qint8)
    print("Task 3/3 Completed!")

#threading
t1 = threading.Thread(target=load_tokenizer)
t2 = threading.Thread(target=load_model)

t1.start()
t2.start()
t1.join()
t2.join()
load_model_quantized()

print("All Threads Done!")



#Add in the text output in a box/scrollable format


#add in open summarised file button
def  open_summary():
    print("Open Summary")
    os.startfile(output_path + "/" + SavingFileName + ".docx")

def reset_fields(self):
    #clear output path
    clear_output_path()
    global outputPathText
    #clear text
    text = ""
    #Clear Saving Name text box
    #saving_name.delete(0, tk.END)
    saving_name = ""
    self.saving_text.delete(0, tk.END)
    
    
    #Clear saving name variable 
    SavingFileName = ""
    #Clear "document Loadings" text box
    DocLoaded.pack_forget()
    
    # Hide Summarise button
    summarise.pack_forget()
    
    #Clear Case file location
    filePath= ""
        
    #Clear Case summarised text
    case_summarised.pack_forget()
    
    #Clear progress bar
    progress_bar.pack_forget()
    
    #Clear 'Summary" Text
    loadingSummaryText.pack_forget()
    #Clear "Summarising Case" text
    
    #clear open summary button
    open_summary_button.pack_forget()
    
    
    
    


class Root(ttk.Window):
    def __init__(self):
        super().__init__()
        
        style = ttk.Style("superhero")
        print("Initialized")
        
        self.geometry("800x600")
       
                    
        #Saving Location Path
        self.folder_selector = ttk.Button(self, text="Select Saving Location", command=SavingLocation)
        self.folder_selector.pack(pady=20)
        
        #Saving Name text
        savingHowTotext = "Please end the name of the summarised File" 
        savingHowTo = ttk.Label(master=self, text=savingHowTotext)
        savingHowTo.pack(pady=10, padx=10)
        
        #saving name
        global saving_name
        global saving_text
        saving_text = tk.StringVar()
        saving_name = "Summary"       
        saving_name = ttk.Entry(self, textvariable=saving_text)
        saving_name.pack(pady=10)
        ToolTip(saving_name, text="Enter the name of the summarised file")
        
        
        save_Button = ttk.Button(self, text="Save File Name", command = savedfileName)
        save_Button.pack(pady=10)
        
        #Case File Selector
        self.file_Selector = ttk.Button(self, text="Select Case File", command=file_path)
        self.file_Selector.pack(pady=10)
        ToolTip(self.file_Selector, text="Select the case file to be summarised")
        
        reset_button = ttk.Button(self, text="Reset Fields", command = reset_fields)
        reset_button.pack(pady=10)
        
        
        
       


root = Root() 

outputPathText = ttk.Label(root) 

# Run app

root.mainloop()